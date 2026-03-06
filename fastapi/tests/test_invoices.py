from datetime import date
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document

from app.services.invoices import (
    _next_invoice_number,
    _next_month,
    _notify_slack,
    _update_dates,
    _update_invoice_number,
    generate_invoice,
)


class TestNextMonth:
    def test_regular_month(self):
        assert _next_month(date(2026, 3, 1)) == date(2026, 4, 1)

    def test_december_wraps_to_january(self):
        assert _next_month(date(2026, 12, 1)) == date(2027, 1, 1)

    def test_january(self):
        assert _next_month(date(2026, 1, 1)) == date(2026, 2, 1)


class TestUpdateDates:
    def _make_doc_with_dates(self, date_text: str, due_date_text: str) -> Document:
        doc = Document()
        p_date = doc.add_paragraph()
        p_date.add_run("Date: ")
        p_date.add_run(date_text)
        p_date.add_run("")

        p_due = doc.add_paragraph()
        p_due.add_run("Due Date: ")
        p_due.add_run(due_date_text)
        p_due.add_run("")

        return doc

    def test_updates_date_fields(self):
        doc = self._make_doc_with_dates("2026-1-05", "2026-2-05")
        _update_dates(doc, date(2026, 3, 1), date(2026, 4, 1))

        date_para = next(p for p in doc.paragraphs if p.text.startswith("Date:"))
        due_para = next(p for p in doc.paragraphs if p.text.startswith("Due Date:"))

        assert "2026-03-01" in date_para.text
        assert "2026-04-01" in due_para.text

    def test_clears_extra_runs(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Due Date: ")
        p.add_run("2026-4-0")
        p.add_run("")
        p.add_run("5")

        _update_dates(doc, date(2026, 3, 1), date(2026, 4, 1))
        assert p.runs[1].text == "2026-04-01"
        assert p.runs[2].text == ""
        assert p.runs[3].text == ""


class TestInvoiceNumber:
    def test_next_number_empty_dir(self, tmp_path):
        assert _next_invoice_number(tmp_path) == "INV000001"

    def test_next_number_with_existing_pdfs(self, tmp_path):
        (tmp_path / "leonel-marinaro_2026-01.pdf").write_bytes(b"fake")
        (tmp_path / "leonel-marinaro_2026-02.pdf").write_bytes(b"fake")
        assert _next_invoice_number(tmp_path) == "INV000003"

    def test_ignores_non_matching_files(self, tmp_path):
        (tmp_path / "leonel-marinaro_2026-01.pdf").write_bytes(b"fake")
        (tmp_path / "other-file.pdf").write_bytes(b"fake")
        assert _next_invoice_number(tmp_path) == "INV000002"

    def test_update_invoice_number_in_doc(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Invoice No: INV000002")
        p.add_run("")

        _update_invoice_number(doc, "INV000005")

        para = next(p for p in doc.paragraphs if "Invoice No" in p.text)
        assert para.text == "Invoice No: INV000005"


class TestGenerateInvoice:
    @pytest.mark.asyncio
    async def test_missing_template_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError, match="Plantilla no encontrada"):
            await generate_invoice({
                "template_path": str(tmp_path / "no_existe.docx"),
                "output_dir": str(tmp_path),
            })

    @pytest.mark.asyncio
    async def test_generates_pdf(self, tmp_path):
        # Crear plantilla de prueba
        doc = Document()
        p_date = doc.add_paragraph()
        p_date.add_run("Date: ")
        p_date.add_run("2026-1-05")

        p_due = doc.add_paragraph()
        p_due.add_run("Due Date: ")
        p_due.add_run("2026-2-05")

        template_path = tmp_path / "template.docx"
        doc.save(str(template_path))

        fake_pdf = tmp_path / "convert" / "invoice.pdf"
        fake_pdf.parent.mkdir()
        fake_pdf.write_bytes(b"%PDF-1.4 fake")

        def mock_run(cmd, **kwargs):
            # Simular que soffice genera el PDF en el outdir
            outdir = Path(cmd[cmd.index("--outdir") + 1])
            docx_name = Path(cmd[-1]).stem
            (outdir / f"{docx_name}.pdf").write_bytes(b"%PDF-1.4 fake")
            return MagicMock(returncode=0, stderr="")

        with patch("app.services.invoices._find_soffice", return_value="soffice"), \
             patch("subprocess.run", side_effect=mock_run):
            result = await generate_invoice({
                "template_path": str(template_path),
                "output_dir": str(tmp_path / "output"),
                "date": "2026-03-01",
            })

        assert result["invoice_number"] == "INV000001"
        assert result["invoice_date"] == "2026-03-01"
        assert result["due_date"] == "2026-04-01"
        assert result["pdf_path"].endswith("leonel-marinaro_2026-03.pdf")
        assert Path(result["pdf_path"]).exists()

    @pytest.mark.asyncio
    async def test_default_date_is_first_of_month(self, tmp_path):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Date: ")
        p.add_run("old")

        template_path = tmp_path / "template.docx"
        doc.save(str(template_path))

        with patch("app.services.invoices._find_soffice", return_value="soffice"), \
             patch("subprocess.run") as mock_run, \
             patch("app.services.invoices.date") as mock_date:
            mock_date.today.return_value = date(2026, 5, 15)
            mock_date.fromisoformat = date.fromisoformat
            mock_date.side_effect = lambda *a, **kw: date(*a, **kw)

            mock_run.return_value = MagicMock(returncode=0, stderr="")
            # Simular creacion del PDF
            def create_fake_pdf(cmd, **kwargs):
                outdir = Path(cmd[cmd.index("--outdir") + 1])
                docx_name = Path(cmd[-1]).stem
                (outdir / f"{docx_name}.pdf").write_bytes(b"%PDF-1.4 fake")
                return MagicMock(returncode=0, stderr="")

            mock_run.side_effect = create_fake_pdf

            result = await generate_invoice({
                "template_path": str(template_path),
                "output_dir": str(tmp_path / "output"),
            })

        assert result["invoice_date"] == "2026-05-01"
        assert result["due_date"] == "2026-06-01"


class TestSlackNotification:
    @pytest.mark.asyncio
    async def test_sends_slack_message(self):
        result = {
            "pdf_path": "/tmp/leonel-marinaro_2026-03.pdf",
            "invoice_number": "INV000003",
            "invoice_date": "2026-03-01",
            "due_date": "2026-04-01",
        }
        mock_response = MagicMock(status_code=200, text="ok")
        mock_client = AsyncMock()
        mock_client.post.return_value = mock_response
        mock_client.__aenter__ = AsyncMock(return_value=mock_client)
        mock_client.__aexit__ = AsyncMock(return_value=False)

        with patch("httpx.AsyncClient", return_value=mock_client):
            await _notify_slack("https://hooks.slack.com/test", result)

        mock_client.post.assert_called_once()
        call_kwargs = mock_client.post.call_args
        assert "Factura generada" in call_kwargs.kwargs["json"]["text"]

    @pytest.mark.asyncio
    async def test_generate_invoice_calls_slack_when_env_set(self, tmp_path):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Date: ")
        p.add_run("old")
        template_path = tmp_path / "template.docx"
        doc.save(str(template_path))

        def create_fake_pdf(cmd, **kwargs):
            outdir = Path(cmd[cmd.index("--outdir") + 1])
            docx_name = Path(cmd[-1]).stem
            (outdir / f"{docx_name}.pdf").write_bytes(b"%PDF-1.4 fake")
            return MagicMock(returncode=0, stderr="")

        with patch("app.services.invoices._find_soffice", return_value="soffice"), \
             patch("subprocess.run", side_effect=create_fake_pdf), \
             patch.dict("os.environ", {"SLACK_WEBHOOK_URL": "https://hooks.slack.com/test"}), \
             patch("app.services.invoices._notify_slack", new_callable=AsyncMock) as mock_slack:
            await generate_invoice({
                "template_path": str(template_path),
                "output_dir": str(tmp_path / "output"),
                "date": "2026-03-01",
            })

        mock_slack.assert_called_once()

    @pytest.mark.asyncio
    async def test_no_slack_when_env_missing(self, tmp_path):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Date: ")
        p.add_run("old")
        template_path = tmp_path / "template.docx"
        doc.save(str(template_path))

        def create_fake_pdf(cmd, **kwargs):
            outdir = Path(cmd[cmd.index("--outdir") + 1])
            docx_name = Path(cmd[-1]).stem
            (outdir / f"{docx_name}.pdf").write_bytes(b"%PDF-1.4 fake")
            return MagicMock(returncode=0, stderr="")

        with patch("app.services.invoices._find_soffice", return_value="soffice"), \
             patch("subprocess.run", side_effect=create_fake_pdf), \
             patch.dict("os.environ", {"SLACK_WEBHOOK_URL": ""}, clear=False), \
             patch("app.services.invoices._notify_slack", new_callable=AsyncMock) as mock_slack:
            await generate_invoice({
                "template_path": str(template_path),
                "output_dir": str(tmp_path / "output"),
                "date": "2026-03-01",
            })

        mock_slack.assert_not_called()
