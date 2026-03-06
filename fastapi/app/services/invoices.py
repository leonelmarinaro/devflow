"""Generación automática de facturas mensuales desde plantilla .docx."""

import os
import shutil
import subprocess
from datetime import date
from pathlib import Path
from tempfile import TemporaryDirectory

import httpx
from docx import Document
from loguru import logger

TEMPLATE_PATH = Path(
    "/Users/lmarinaro/Documents/Leo/Facturas/"
    "EXT - MAKE A COPY - Modelo Factura Contractor GKT.docx"
)
OUTPUT_DIR = Path("/Users/lmarinaro/Documents/Leo/Facturas")

SOFFICE_PATHS = [
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "soffice",
]


def _find_soffice() -> str:
    for path in SOFFICE_PATHS:
        if Path(path).exists():
            return path
    result = subprocess.run(["which", "soffice"], capture_output=True, text=True)
    if result.returncode == 0:
        return result.stdout.strip()
    raise FileNotFoundError(
        "LibreOffice no encontrado. Instalalo con: brew install --cask libreoffice"
    )


def _next_month(d: date) -> date:
    if d.month == 12:
        return d.replace(year=d.year + 1, month=1)
    return d.replace(month=d.month + 1)


def _next_invoice_number(output_dir: Path) -> str:
    existing = list(output_dir.glob("leonel-marinaro_*.pdf"))
    return f"INV{len(existing) + 1:06d}"


def _update_invoice_number(doc: Document, invoice_number: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Invoice No:"):
            paragraph.runs[0].text = f"Invoice No: {invoice_number}"
            for run in paragraph.runs[1:]:
                run.text = ""
            break


def _update_dates(doc: Document, invoice_date: date, due_date: date) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.startswith("Date:") and not text.startswith("Due Date:"):
            if len(paragraph.runs) >= 2:
                paragraph.runs[1].text = str(invoice_date)
                for run in paragraph.runs[2:]:
                    run.text = ""
        elif text.startswith("Due Date:") and len(paragraph.runs) >= 2:
            paragraph.runs[1].text = str(due_date)
            for run in paragraph.runs[2:]:
                run.text = ""


def _convert_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    soffice = _find_soffice()
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir",
         str(output_dir), str(docx_path)],
        capture_output=True,
        text=True,
        timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(f"Error convirtiendo a PDF: {result.stderr}")
    return output_dir / docx_path.with_suffix(".pdf").name


async def generate_invoice(payload: dict) -> dict:
    """
    Genera factura PDF desde la plantilla .docx.

    Payload opcional:
    - template_path: ruta a la plantilla .docx
    - output_dir: directorio de salida
    - date: fecha de la factura YYYY-MM-DD (default: 1ro del mes actual)
    """
    template = Path(payload.get("template_path", str(TEMPLATE_PATH)))
    output_dir = Path(payload.get("output_dir", str(OUTPUT_DIR)))

    if payload.get("date"):
        invoice_date = date.fromisoformat(payload["date"])
    else:
        invoice_date = date.today().replace(day=1)

    due_date = _next_month(invoice_date)

    if not template.exists():
        raise FileNotFoundError(f"Plantilla no encontrada: {template}")

    output_dir.mkdir(parents=True, exist_ok=True)
    output_filename = f"leonel-marinaro_{invoice_date.strftime('%Y-%m')}.pdf"
    output_path = output_dir / output_filename

    logger.info(f"Generando factura: {output_filename} (Date: {invoice_date}, Due: {due_date})")

    invoice_number = _next_invoice_number(output_dir)

    doc = Document(str(template))
    _update_dates(doc, invoice_date, due_date)
    _update_invoice_number(doc, invoice_number)

    with TemporaryDirectory() as tmpdir:
        tmp_docx = Path(tmpdir) / "invoice.docx"
        doc.save(str(tmp_docx))

        pdf_path = _convert_to_pdf(tmp_docx, Path(tmpdir))
        shutil.move(str(pdf_path), str(output_path))

    logger.info(f"Factura generada: {output_path}")

    result = {
        "pdf_path": str(output_path),
        "invoice_number": invoice_number,
        "invoice_date": str(invoice_date),
        "due_date": str(due_date),
    }

    slack_url = os.environ.get("SLACK_WEBHOOK_URL", "")
    if slack_url:
        await _notify_slack(slack_url, result)

    return result


async def _notify_slack(webhook_url: str, result: dict) -> None:
    text = (
        f":page_facing_up: *Factura generada*\n"
        f"- N°: `{result['invoice_number']}`\n"
        f"- Periodo: `{result['invoice_date']}` → `{result['due_date']}`\n"
        f"- Archivo: `{Path(result['pdf_path']).name}`"
    )
    try:
        async with httpx.AsyncClient() as client:
            resp = await client.post(webhook_url, json={"text": text}, timeout=10)
            if resp.status_code == 200:
                logger.info("Notificación Slack enviada")
            else:
                logger.warning(f"Slack respondió {resp.status_code}: {resp.text[:200]}")
    except httpx.HTTPError as e:
        logger.error(f"Error enviando a Slack: {e}")
